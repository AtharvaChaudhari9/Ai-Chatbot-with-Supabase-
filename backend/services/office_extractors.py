import io
import csv
import logging
import subprocess
import tempfile
import os

logger = logging.getLogger("rag_backend.office_extractors")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text paragraphs and table rows from docx file bytes.
    """
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed.")
        raise RuntimeError("Word document parser (python-docx) is not installed in the environment.")

    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []

    # 1. Extract standard text paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # 2. Extract table contents (separating cells with pipe delimiters)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if any(row_text):
                full_text.append(" | ".join(row_text))

    return "\n".join(full_text)


def extract_text_from_doc(file_bytes: bytes) -> str:
    """
    Extracts text from a .doc binary file using antiword command-line utility.
    """
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name
        
        # Run antiword
        result = subprocess.run(["antiword", "-m", "UTF-8", tmp_path], capture_output=True, text=True, check=True)
        text = result.stdout.strip()
        if not text:
            # Try without formatting flags
            result = subprocess.run(["antiword", tmp_path], capture_output=True, text=True, check=True)
            text = result.stdout.strip()
        return text
    except Exception as e:
        logger.error(f"antiword failed to parse .doc file: {e}")
        raise RuntimeError(f"Failed to extract text from Word .doc file. Make sure antiword is installed. Error: {str(e)}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """
    Extracts grid contents from all sheets of an xlsx spreadsheet.
    """
    try:
        import openpyxl
    except ImportError:
        logger.error("openpyxl is not installed.")
        raise RuntimeError("Excel document parser (openpyxl) is not installed in the environment.")

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    full_text = []

    for sheet_name in wb.sheetnames:
        full_text.append(f"--- Sheet: {sheet_name} ---")
        ws = wb[sheet_name]
        
        for row in ws.iter_rows(values_only=True):
            # Format spreadsheet row as a pipe-separated line
            row_vals = [str(val).strip() for val in row if val is not None]
            if any(row_vals):
                full_text.append(" | ".join(row_vals))

    return "\n".join(full_text)


def extract_text_from_xls(file_bytes: bytes) -> str:
    """
    Extracts grid contents from all sheets of an xls spreadsheet using xlrd.
    """
    try:
        import xlrd
    except ImportError:
        logger.error("xlrd is not installed.")
        raise RuntimeError("Excel .xls document parser (xlrd) is not installed in the environment.")

    try:
        workbook = xlrd.open_workbook(file_contents=file_bytes)
        full_text = []

        for sheet_idx in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(sheet_idx)
            full_text.append(f"--- Sheet: {sheet.name} ---")
            
            for row_idx in range(sheet.nrows):
                row_vals = []
                for col_idx in range(sheet.ncols):
                    val = sheet.cell_value(row_idx, col_idx)
                    if val is not None and str(val).strip() != "":
                        row_vals.append(str(val).strip())
                if any(row_vals):
                    full_text.append(" | ".join(row_vals))

        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"xlrd failed to parse xls: {e}")
        raise RuntimeError(f"Failed to parse Excel .xls file: {str(e)}")


def extract_text_from_csv(file_bytes: bytes) -> str:
    """
    Extracts contents from a CSV sheet using Python's built-in csv module.
    """
    content = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(content))
    full_text = []

    for row in reader:
        row_vals = [cell.strip() for cell in row if cell.strip()]
        if any(row_vals):
            full_text.append(" | ".join(row_vals))

    return "\n".join(full_text)
